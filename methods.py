# methods.py
import pandas as pd

import os
from werkzeug.utils import secure_filename
import fitz
import sys
from docx import Document
ALLOWED_EXTENSIONS = {'pdf', 'txt','docx','doc','csv','xlsx'}

def is_allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_text_from_pdf(pdf_path):
    text = ""
    with fitz.open(pdf_path) as doc:
        for page_number in range(doc.page_count):
            page = doc[page_number]
            text += page.get_text()

    return text
if sys.platform == 'win32':
    sys.stdout.reconfigure(encoding='utf-8')
def extract_text_from_csv(csv_path):
    try:
        # Read CSV file into a pandas DataFrame
        df = pd.read_csv(csv_path,encoding='latin-1')
        
        # Extract text from DataFrame
        text_content = df.to_string(index=False)
        
        return text_content
    except Exception as e:
        return f"Error extracting text from CSV: {str(e)}"
def extract_text_from_excel(excel_path, sheet_name=None):
    try:
        # Read Excel file into a pandas DataFrame using 'openpyxl'
        df = pd.read_excel(excel_path, sheet_name=sheet_name)
        
        
            # Extract text from DataFrame
        text_content = str(df)
        return text_content
        
    except Exception as e:
        return f"Error extracting text from Excel: {str(e)}"
def extract_text_from_text_file(file_path):
    try:
        with open(file_path, 'r', encoding='utf-8') as file:
            text = file.read()
        return text
    except Exception as e:
        return f"Error extracting text from text file: {str(e)}"
from docx import Document
import aspose.words as aw

def extract_text_from_docx(docx_path):
    try:
        doc = Document(docx_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + '\n'
        return text
    except Exception as e:
        return f"Error extracting text from docx file: {str(e)}"
def convert_doc_to_text(input_doc_path):
    try:
        # Load the Word document using Aspose.Words
        doc = aw.Document(input_doc_path)

        # Extract the text from the document
        text_content = doc.get_text()

        return text_content
    except Exception as e:
        return f"Error converting doc to text: {str(e)}"
def handle_file_upload(files):
    concatenated_text = ""

    for file in files:
        # Check if the file is allowed
        if not file or not is_allowed_file(file.filename):
            return None, f"Invalid file type for file {file.filename}. Allowed file types: pdf, txt"

        # Generate a secure filename and save the file
        filename = secure_filename(file.filename)
        file_path = os.path.join("uploads", filename)
        file.save(file_path)

        # Extract text based on file type
        if filename.lower().endswith('.pdf'):
            text = extract_text_from_pdf(file_path)
        elif filename.lower().endswith('.txt'):
            text = extract_text_from_text_file(file_path)
        elif filename.lower().endswith('.docx'):
            text = extract_text_from_docx(file_path)
        elif filename.lower().endswith('.doc'):
            text = convert_doc_to_text(file_path) 
        elif filename.lower().endswith('.csv'):
            text = extract_text_from_csv(file_path) 
        elif filename.lower().endswith('.xlsx'):
            text = extract_text_from_excel(file_path)      
        else:
            return None, f"Unsupported file type for file {file.filename}"

        # Concatenate the extracted text
        concatenated_text += text + "\n"

        # Clean up: delete the uploaded file
        os.remove(file_path)

    return concatenated_text, None
